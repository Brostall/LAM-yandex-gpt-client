import os
import re
from datetime import datetime
from docx import Document
import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment
from typing import Optional, Dict, Any, List, Tuple, Set
import json
import logging
from openpyxl.utils import get_column_letter
from docx.shared import Pt, RGBColor

class FileHandler:
    def __init__(self, team_name: str, base_path: str = "data"):
        """
        Инициализация обработчика файлов
        
        Args:
            team_name (str): название команды
            base_path (str): базовая директория для хранения файлов
        """
        self.team_name = team_name
        self.base_path = base_path
        self.messages_path = os.path.join(base_path, "messages", team_name)
        self.excel_path = os.path.join(base_path, "excel")
        
        # необходимые директории
        os.makedirs(self.base_path, exist_ok=True)
        os.makedirs(self.messages_path, exist_ok=True)
        os.makedirs(self.excel_path, exist_ok=True)
        
        # Счетчик сообщений для каждого отправителя
        self.message_counters: Dict[str, int] = {}
        
        # Статистика обработанных сообщений
        self.statistics = {
            "total_messages": 0,
            "senders": {},
            "fields_data": {}
        }
        
        # Загружаем справочную информацию (будет использоваться для проверки идентификации)
        self.reference_data = self.load_reference_data()
        
        # Текущий Excel-файл (используется для обновления после каждого сообщения)
        self.current_excel_path = None
    
    def load_reference_data(self):
        """
        Загружает справочные данные для классификации сообщений.
        Если файл справочных данных не найден, использует значения по умолчанию.
        """
        reference_file = os.path.join(self.base_path, "reference_data.json")
        
        try:
            if os.path.exists(reference_file):
                with open(reference_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                logging.info(f"Загружены справочные данные из {reference_file}")
                return data
            else:
                logging.info("Файл справочных данных не найден, используются значения по умолчанию")
                return self._get_default_reference_data()
        except Exception as e:
            logging.warning(f"Не удалось загрузить справочные данные: {str(e)}")
            return self._get_default_reference_data()
    
    def _get_default_reference_data(self):
        """
        Возвращает значения справочных данных по умолчанию.
        """
        return {
            "subdivisions": [
                "Центральное отделение", 
                "Северное отделение", 
                "Южное отделение", 
                "Восточное отделение", 
                "Западное отделение",
                "АО Кроноткинское",
                "Восход",
                "Колхоз Прогресс",
                "Мир",
                "СП Коломейцево"
            ],
            "operations": [
                "Вспашка", 
                "Боронование", 
                "Культивация",
                "Культивация предпосевная",
                "Культивация сплошная",
                "Дискование",
                "Гербицидная обработка",
                "Гербицидная обработка сплошная",
                "Фунгицидная обработка", 
                "Инсектицидная обработка", 
                "Внесение удобрений", 
                "Посев", 
                "Уборка", 
                "Кошение",
                "Скашивание",
                "Орошение",
                "Опрыскивание",
                "Сбор урожая",
                "Транспортировка",
                "Погрузка",
                "Разгрузка",
                "Плановое ТО",
                "Внеплановый ремонт",
                "Перегон техники"
            ],
            "cultures": [
                "Вика+Тритикале", "Горох на зерно", "Горох товарный", "Гуар", "Конопля", 
                "Кориандр", "Кукуруза кормовая", "Кукуруза семенная", "Кукуруза товарная", 
                "Люцерна", "Многолетние злаковые травы", "Многолетние травы прошлых лет", 
                "Многолетние травы текущего года", "Овес", "Подсолнечник кондитерский", 
                "Подсолнечник семенной", "Подсолнечник товарный", "Просо", 
                "Пшеница озимая на зеленый корм", "Пшеница озимая семенная", 
                "Пшеница озимая товарная", "Рапс озимый", "Рапс яровой", "Свекла сахарная", 
                "Сорго", "Сорго кормовой", "Сорго-суданковый гибрид", "Соя семенная", 
                "Соя товарная", "Чистый пар", "Чумиза", "Ячмень озимый", "Ячмень озимый семенной"
            ],
            "pu_subdivisions": {
                "АОР": ["Кавказ", "Север", "Центр", "Юг", "Рассвет"],
                "ТСК": ["Нет ПУ"],
                "АО Кроноткинское": ["Нет ПУ"],
                "Восход": ["Нет ПУ"],
                "Колхоз Прогресс": ["Нет ПУ"],
                "Мир": ["Нет ПУ"],
                "СП Коломейцево": ["Нет ПУ"]
            },
            "departments": {
                "Кавказ": ["18", "19"],
                "Север": ["3", "7", "10", "20"],
                "Центр": ["1", "4", "5", "6", "9"],
                "Юг": ["11", "12", "16", "17"],
                "Рассвет": []
            },
            "operations_by_culture": {
                "Пшеница озимая товарная": [
                    "Пахота", "Дискование", "Выравнивание", "Предпосевная культивация", 
                    "Сев", "Подкормка", "2-я подкормка", "Гербицидная обработка",
                    "Фунгицидная обработка", "Уборка", "Прикатывание посевов"
                ],
                "Соя товарная": [
                    "Пахота", "Дискование", "Предпосевная культивация", "Сев", 
                    "Гербицидная обработка", "Междурядная обработка", "Уборка"
                ],
                "Подсолнечник товарный": [
                    "Пахота", "Дискование", "Предпосевная культивация", "Сев", 
                    "Гербицидная обработка", "Междурядная обработка", "Уборка"
                ],
                "Свекла сахарная": [
                    "Пахота", "Дискование", "Предпосевная культивация", "Сев", 
                    "Гербицидная обработка", "Междурядная обработка", "Уборка"
                ],
                "Многолетние травы": [
                    "Пахота", "Дискование", "Предпосевная культивация", "Сев", 
                    "Уборка", "Подкормка"
                ]
            },
            "culture_abbreviations": {
                "мн тр": "Многолетние травы прошлых лет",
                "мн тр тек.года": "Многолетние травы текущего года",
                "мн зл": "Многолетние злаковые травы",
                "оз пш": "Пшеница озимая товарная",
                "оз пш сем": "Пшеница озимая семенная",
                "оз пш на зел": "Пшеница озимая на зеленый корм",
                "сах св": "Свекла сахарная",
                "соя": "Соя товарная",
                "соя сем": "Соя семенная",
                "подс": "Подсолнечник товарный",
                "подс сем": "Подсолнечник семенной",
                "подс кон": "Подсолнечник кондитерский",
                "яч": "Ячмень озимый", 
                "оз яч": "Ячмень озимый",
                "оз яч сем": "Ячмень озимый семенной",
                "кук": "Кукуруза товарная",
                "кук корм": "Кукуруза кормовая",
                "кук сем": "Кукуруза семенная",
                "рапс": "Рапс яровой",
                "оз рапс": "Рапс озимый",
                "горох": "Горох на зерно",
                "горох тов": "Горох товарный",
                "вика+трит": "Вика+Тритикале",
                "сорго": "Сорго",
                "сорго корм": "Сорго кормовой",
                "гибрид": "Сорго-суданковый гибрид"
            }
        }
    
    async def save_message(self, sender_name: str, message_text: str, yandex_gpt=None) -> Tuple[str, Dict[str, Any]]:
        """
        Сохраняет сообщение в файл Word, анализирует с помощью YandexGPT и обновляет Excel
        
        Args:
            sender_name (str): имя отправителя
            message_text (str): текст сообщения
            yandex_gpt: экземпляр класса YandexGPT
            
        Returns:
            Tuple[str, Dict[str, Any]]: (путь к сохраненному файлу, результат анализа)
        """
        # Увеличиваем счетчик сообщений для отправителя
        if sender_name not in self.message_counters:
            self.message_counters[sender_name] = 0
        self.message_counters[sender_name] += 1
        
        # Обновляем статистику
        self.statistics["total_messages"] += 1
        if sender_name not in self.statistics["senders"]:
            self.statistics["senders"][sender_name] = 0
        self.statistics["senders"][sender_name] += 1
        
        # Формируем имя файла для текущего дня
        now = datetime.now()
        filename = f"messages_{now.strftime('%d%m%Y')}.docx"
        filepath = os.path.join(self.messages_path, filename)
        
        try:
            # Если файл существует, открываем его
            if os.path.exists(filepath):
                doc = Document(filepath)
            else:
                # Создаем новый документ
                doc = Document()
                # Добавляем заголовок
                doc.add_heading(f'Сообщения агрономов за {now.strftime("%d.%m.%Y")}', 0)
            
            # Добавляем информацию о сообщении
            doc.add_heading(f'{sender_name} - {now.strftime("%H:%M")}', level=2)
            doc.add_paragraph(message_text)
            doc.add_paragraph('') # Пустая строка для разделения сообщений
            
            # Сохраняем документ
            doc.save(filepath)
            logging.info(f"Сообщение добавлено в файл: {filepath}")
            
        except Exception as e:
            logging.error(f"Ошибка при сохранении сообщения в Word: {str(e)}")
            # В случае ошибки создаем новый файл
            try:
                doc = Document()
                doc.add_heading(f'Сообщения агрономов за {now.strftime("%d.%m.%Y")}', 0)
                doc.add_heading(f'{sender_name} - {now.strftime("%H:%M")}', level=2)
                doc.add_paragraph(message_text)
                doc.add_paragraph('')
                doc.save(filepath)
                logging.info(f"Создан новый файл с сообщением: {filepath}")
            except Exception as e:
                logging.error(f"Критическая ошибка при создании нового файла Word: {str(e)}")
                raise
        
        # Анализируем сообщение, если предоставлен экземпляр YandexGPT
        analysis_result = {}
        if yandex_gpt:
            try:
                analysis_result = await self.analyze_and_cache_message(sender_name, message_text, yandex_gpt)
            except Exception as e:
                logging.error(f"Ошибка при анализе сообщения: {str(e)}")
                # Используем обычный парсинг при ошибке
                analysis_result = self.parse_message(message_text)
        
        # Автоматически обновляем Excel после каждого сообщения
        try:
            self.update_excel()
        except Exception as e:
            logging.error(f"Ошибка при автоматическом обновлении Excel: {str(e)}")
        
        return filepath, analysis_result
    
    def parse_message(self, message_text: str) -> Dict[str, Any]:
        """
        Анализирует текст сообщения и извлекает нужную информацию
        
        Args:
            message_text (str): текст сообщения
            
        Returns:
            Dict[str, Any]: извлеченные данные
        """
        # Инициализируем базовые поля
        data = {
            "work_type": "",
            "operation": "",
            "culture_from": "",
            "culture_to": "",
            "pu_number": "",
            "pu_area": "",
            "department": "",
            "department_number": "",
            "department_area": "",
            "date": "",
            "subdivision": "АОР",
            "val_day": "",
            "val_total": "",
            "date_processed": datetime.now().strftime("%d.%m.%Y")
        }
        
        try:
            # Разбиваем сообщение на строки и операции
            operations = []
            current_operation = None
            lines = message_text.strip().split('\n')
            
            # Шаблоны регулярных выражений для извлечения данных
            patterns = {
                "field_info": r"(?:Предп|диск|Пахота|Выравн)(?:ов|п)?(?:ание|а|)?\s+(\w+(?:\s+\w+)?)\s+(?:под|на)?\s*(\w+(?:\s+\w+)?)?",
                "pu_info": r"По\s*Пу\s*(\d+)/(\d+)",
                "department_info": r"Отд\s*(\d+)\s*(\d+)/(\d+)",
                "date": r"(\d{1,2}\.\d{1,2}(?:\.\d{4})?)",
                "operation": r"^([А-Яа-я0-9\-\s]+)(?:\n|$)",
                "val_day": r"Вал\s+за\s+день\s+(\d+[\s,\.]\d+)",
                "val_total": r"Вал\s+с\s+начала\s+(\d+[\s,\.]\d+)"
            }
            
            # Обработка заголовка сообщения для поиска даты и подразделения
            for i, line in enumerate(lines):
                date_match = re.search(patterns["date"], line)
                if date_match:
                    date = date_match.group(1)
                    if len(date.split('.')) == 2:
                        date += f".{datetime.now().year}"
                    data["date"] = date
                
                # Поиск подразделения в начале сообщения
                for sub in self.reference_data["subdivisions"]:
                    if sub.lower() in line.lower():
                        data["subdivision"] = sub
                        break
            
            # Режим обработки - операция по операции
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                
                if not line:
                    i += 1
                    continue
                
                # Проверяем, является ли строка операцией
                operation_match = re.search(patterns["operation"], line)
                culture_match = re.search(patterns["field_info"], line)
                
                if operation_match or (culture_match and "под" in line):
                    # Начало новой операции
                    if current_operation:
                        operations.append(current_operation)
                    
                    current_operation = data.copy()
                    
                    # Извлекаем название операции
                    if operation_match:
                        current_operation["work_type"] = operation_match.group(1).strip()
                        current_operation["operation"] = operation_match.group(1).strip()
                    
                    # Проверяем наличие культуры в той же строке
                    if culture_match:
                        current_operation["culture_from"] = culture_match.group(1)
                        if culture_match.group(2):
                            current_operation["culture_to"] = self.normalize_culture_abbreviation(culture_match.group(2))
                
                # Проверяем информацию о ПУ
                pu_match = re.search(patterns["pu_info"], line)
                if pu_match:
                    if current_operation:
                        current_operation["pu_number"] = pu_match.group(1)
                        current_operation["pu_area"] = pu_match.group(2)
                
                # Проверяем информацию об отделах
                dept_match = re.search(patterns["department_info"], line)
                if dept_match:
                    if current_operation:
                        if not current_operation["department"]:
                            current_operation["department"] = dept_match.group(1)
                            current_operation["department_number"] = dept_match.group(2)
                            current_operation["department_area"] = dept_match.group(3)
                        else:
                            # Добавляем через запятую, если уже есть отделы
                            current_operation["department"] += f", {dept_match.group(1)}"
                            current_operation["department_number"] += f", {dept_match.group(2)}"
                            current_operation["department_area"] += f", {dept_match.group(3)}"
                
                # Проверяем информацию о вале
                val_day_match = re.search(patterns["val_day"], line)
                if val_day_match and current_operation:
                    current_operation["val_day"] = val_day_match.group(1).replace(" ", "")
                
                val_total_match = re.search(patterns["val_total"], line)
                if val_total_match and current_operation:
                    current_operation["val_total"] = val_total_match.group(1).replace(" ", "")
                
                i += 1
            
            # Добавляем последнюю операцию, если она есть
            if current_operation:
                operations.append(current_operation)
            
            # Если операции найдены, возвращаем их
            if operations:
                result = {
                    "operations": operations,
                    "date_processed": data["date_processed"]
                }
                
                # Для совместимости с другими методами, добавляем поля первой операции
                if operations:
                    for key, value in operations[0].items():
                        if key != "operations":
                            result[key] = value
                
                return result
            
            # Если операций не найдено, используем простой подход
            operation = data.copy()
            
            # Проверяем наличие операции в первой строке
            if lines and lines[0]:
                operation_match = re.search(r"^([А-Яа-я\s\-]+)", lines[0])
                if operation_match:
                    operation["work_type"] = operation_match.group(1).strip()
                    operation["operation"] = operation_match.group(1).strip()
            
            # Ищем информацию о ПУ и отделах во всем сообщении
            for line in lines:
                pu_match = re.search(patterns["pu_info"], line)
                if pu_match:
                    operation["pu_number"] = pu_match.group(1)
                    operation["pu_area"] = pu_match.group(2)
                
                dept_match = re.search(patterns["department_info"], line)
                if dept_match:
                    if not operation["department"]:
                        operation["department"] = dept_match.group(1)
                        operation["department_number"] = dept_match.group(2)
                        operation["department_area"] = dept_match.group(3)
                    else:
                        operation["department"] += f", {dept_match.group(1)}"
                        operation["department_number"] += f", {dept_match.group(2)}"
                        operation["department_area"] += f", {dept_match.group(3)}"
            
            # Возвращаем результат
            logging.info(f"Успешно разобрано сообщение базовым парсером")
            return operation
            
        except Exception as e:
            logging.error(f"Ошибка при базовом парсинге сообщения: {str(e)}")
            return data
    
    async def analyze_with_yandex_gpt(self, message_text: str, yandex_gpt) -> Dict[str, Any]:
        """
        Анализирует сообщение с помощью YandexGPT с использованием обучения на примерах
        
        Args:
            message_text (str): текст сообщения
            yandex_gpt: экземпляр класса YandexGPT
            
        Returns:
            Dict[str, Any]: извлеченные данные
        """
        # Примеры для обучения - показываем модели, как обрабатывать сообщения
        examples = [
            {
                "message": "Пахота зяби под мн тр\nПо Пу 26/488\nОтд 12 26/221",
                "analysis": {
                    "work_type": "Пахота", 
                    "operation": "Пахота зяби", 
                    "culture_from": "зяби", 
                    "culture_to": "Многолетние", 
                    "pu_number": "26", 
                    "pu_area": "488", 
                    "department": "12", 
                    "department_number": "26", 
                    "department_area": "221",
                    "date": "",
                    "subdivision": "АОР",
                    "val_day": "",
                    "val_total": ""
                }
            },
            {
                "message": "Предп культ под оз пш\nПо Пу 215/1015\nОтд 12 128/317\nОтд 16 123/529",
                "analysis": {
                    "work_type": "Предпосевная культивация",
                    "operation": "Предп культ",
                    "culture_from": "культ",
                    "culture_to": "Пшеница озимая товарная",
                    "pu_number": "215",
                    "pu_area": "1015",
                    "department": "12, 16",
                    "department_number": "128, 123",
                    "department_area": "317, 529",
                    "date": "",
                    "subdivision": "АОР",
                    "val_day": "",
                    "val_total": ""
                }
            },
            {
                "message": "10.03 день\n2-я подкормка\nПо Пу 1749/2559",
                "analysis": {
                    "work_type": "2-я подкормка",
                    "operation": "2-я подкормка",
                    "culture_from": "",
                    "culture_to": "Пшеница озимая товарная",
                    "pu_number": "1749",
                    "pu_area": "2559",
                    "department": "",
                    "department_number": "",
                    "department_area": "",
                    "date": "10.03.2024",
                    "subdivision": "АОР",
                    "val_day": "",
                    "val_total": ""
                }
            }
        ]
        
        # Формируем список справочных данных для модели
        reference_data_str = json.dumps({
            "subdivisions": list(self.reference_data["subdivisions"]),
            "operations": list(self.reference_data["operations"]),
            "cultures": list(self.reference_data["cultures"])
        }, ensure_ascii=False)
        
        # Формируем примеры для обучения
        examples_str = "\n\n".join([
            f"СООБЩЕНИЕ: {ex['message']}\nАНАЛИЗ: {json.dumps(ex['analysis'], ensure_ascii=False)}"
            for ex in examples[:2]  # Используем только первые два примера для сокращения запроса
        ])
        
        # Улучшенный промпт с примерами и справочными данными
        prompt = f"""
        Ты - специализированная система для анализа сельскохозяйственных сообщений от агрономов.
        
        СПРАВОЧНЫЕ ДАННЫЕ:
        {reference_data_str}
        
        ПРИМЕРЫ АНАЛИЗА:
        {examples_str}
        
        Теперь проанализируй следующее сообщение от агронома и извлеки из него все структурированные данные.
        Сообщение может содержать одну или несколько операций. Каждая операция обычно включает:
        1. Тип работы (Пахота, Дискование, Предпосевная культивация и т.д.)
        2. Данные ПУ (производственный участок) в формате "По Пу номер/площадь"
        3. Данные отделов в формате "Отд номер площадь/площадь"
        
        СООБЩЕНИЕ ДЛЯ АНАЛИЗА:
        {message_text}
        
        Верни JSON с операциями.
        """
        
        try:
            # Вызываем YandexGPT с улучшенным промптом
            logging.info(f"Отправляем запрос на анализ сообщения в YandexGPT, длина сообщения: {len(message_text)}")
            result = await yandex_gpt.generate_response(
                prompt=prompt,
                model="lite",  # Используем lite модель для уменьшения времени обработки
                temperature=0.1,  # Низкая температура для более детерминированных результатов
                system_prompt="Анализируй сельскохозяйственные сообщения и извлекай структурированные данные."
            )
            
            if "error" in result:
                logging.error(f"Ошибка при запросе к YandexGPT: {result['error']}")
                return self.parse_message(message_text)
            
            response_text = yandex_gpt.get_response_text(result)
            
            # Если ответ содержит ошибку, используем обычный парсинг
            if response_text.startswith("Произошла ошибка") or response_text.startswith("Не удалось") or response_text.startswith("Ошибка"):
                logging.warning(f"YandexGPT вернул ошибку: {response_text}")
                return self.parse_message(message_text)
            
            # Извлекаем JSON из ответа
            json_str = response_text.strip()
            
            # Чистим ответ для получения валидного JSON
            # Очищаем строку от маркеров кода, если они есть
            if "```json" in json_str:
                json_str = json_str.split("```json")[1].split("```")[0].strip()
            elif "```" in json_str:
                json_str = json_str.split("```")[1].strip()
                
            # Удаляем все, что не является JSON
            if not json_str.startswith('{'):
                start_idx = json_str.find('{')
                if start_idx == -1:
                    logging.error("Не удалось найти JSON в ответе")
                    return self.parse_message(message_text)
                json_str = json_str[start_idx:]
            
            if not json_str.endswith('}'):
                end_idx = json_str.rfind('}')
                if end_idx == -1:
                    logging.error("Не удалось найти конец JSON в ответе")
                    return self.parse_message(message_text)
                json_str = json_str[:end_idx+1]
            
            try:
                # Парсим JSON
                data = json.loads(json_str)
                logging.info(f"Успешно получен и разобран JSON ответ от YandexGPT")
            except json.JSONDecodeError as e:
                logging.error(f"Ошибка при парсинге JSON: {e}. JSON: {json_str[:200]}...")
                return self.parse_message(message_text)
            
            # Извлекаем первую операцию, если это массив операций
            if "operations" in data and len(data["operations"]) > 0:
                first_operation = data["operations"][0]
                
                # Добавляем остальные поля из первой операции
                result_data = {
                    "operations": data["operations"],
                    "work_type": first_operation.get("work_type", ""),
                    "culture_from": first_operation.get("culture_from", ""),
                    "culture_to": first_operation.get("culture_to", ""),
                    "pu_number": first_operation.get("pu_number", ""),
                    "pu_area": first_operation.get("pu_area", ""),
                    "department": first_operation.get("department", ""),
                    "department_number": first_operation.get("department_number", ""),
                    "department_area": first_operation.get("department_area", ""),
                    "date": first_operation.get("date", ""),
                    "subdivision": first_operation.get("subdivision", "АОР"),
                    "val_day": first_operation.get("val_day", ""),
                    "val_total": first_operation.get("val_total", "")
                }
                
                # Добавляем исправления, если они есть
                if "corrections" in data and data["corrections"]:
                    result_data["corrections"] = data["corrections"]
                
                # Добавляем дату обработки, если нет даты
                if not result_data["date"]:
                    result_data["date_processed"] = datetime.now().strftime("%d.%m.%Y")
                
                return result_data
            else:
                # Если операций нет, заполняем базовые поля
                basic_data = {
                    "date_processed": datetime.now().strftime("%d.%m.%Y"),
                    "operations": []
                }
                
                # Объединяем с полученными данными, если они есть
                if isinstance(data, dict):
                    basic_data.update(data)
                    
                return basic_data
                
        except Exception as e:
            logging.error(f"Ошибка при анализе сообщения с YandexGPT: {str(e)}")
            # Если что-то пошло не так, используем обычный парсинг
            return self.parse_message(message_text)
    
    def get_all_messages(self) -> List[Tuple[str, str, str]]:
        """
        Получает все сохраненные сообщения
        
        Returns:
            List[Tuple[str, str, str]]: список кортежей (имя_отправителя, дата, текст_сообщения)
        """
        messages = []
        
        if not os.path.exists(self.messages_path):
            return messages
            
        for filename in os.listdir(self.messages_path):
            if filename.endswith(".docx"):
                try:
                    # Извлекаем имя отправителя и дату из имени файла
                    parts = filename.split("_")
                    sender_name = parts[0]
                    if len(parts) > 2:
                        date_str = parts[2].replace(".docx", "")
                        # Преобразуем формат даты из МинутаЧасДеньМесяцГод
                        try:
                            date_obj = datetime.strptime(date_str, "%M%H%d%m%Y")
                            date_str = date_obj.strftime("%d.%m.%Y")
                        except:
                            date_str = datetime.now().strftime("%d.%m.%Y")
                    else:
                        date_str = datetime.now().strftime("%d.%m.%Y")
                    
                    # Читаем текст из файла
                    filepath = os.path.join(self.messages_path, filename)
                    doc = Document(filepath)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    
                    messages.append((sender_name, date_str, text))
                except Exception as e:
                    print(f"Ошибка при чтении файла {filename}: {e}")
        
        return messages
    
    def update_excel(self, template_path: str = None) -> str:
        """
        Обновляет Excel файл с данными из всех сообщений
        
        Args:
            template_path (str, optional): путь к шаблону Excel
            
        Returns:
            str: путь к сохраненному файлу
        """
        # Создаем новую книгу Excel каждый раз
        wb = openpyxl.Workbook()
        # Удаляем стандартный лист
        std_sheet = wb.active
        wb.remove(std_sheet)
        # Создаем новый лист с именем
        ws = wb.create_sheet("Отчет агрономов")
        
        # Задаем стили для ячеек
        header_fill = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")
        date_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
        subdivision_fill = PatternFill(start_color="5B9BD5", end_color="5B9BD5", fill_type="solid") 
        operation_fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")
        culture_fill = PatternFill(start_color="9933FF", end_color="9933FF", fill_type="solid")
        day_area_fill = PatternFill(start_color="FF9966", end_color="FF9966", fill_type="solid")
        total_area_fill = PatternFill(start_color="00FF00", end_color="00FF00", fill_type="solid")
        day_val_fill = PatternFill(start_color="FF99CC", end_color="FF99CC", fill_type="solid")
        total_val_fill = PatternFill(start_color="99FF99", end_color="99FF99", fill_type="solid")
        
        # Добавляем раздел легенда в начало файла
        ws.cell(row=1, column=2).value = "Легенда"
        ws.merge_cells('B1:I1')
        
        ws.cell(row=2, column=1).value = "Цветовое обозначение"
        ws.merge_cells('A2:A3')
        
        # Добавляем заголовки колонок с цветами
        headers = ["Дата", "Подразделение", "Операция", "Культура", "За день, га", "С начала операции, га", "Вал за день, ц", "Вал с начала, ц"]
        fills = [date_fill, subdivision_fill, operation_fill, culture_fill, day_area_fill, total_area_fill, day_val_fill, total_val_fill]
        
        for col, (header, fill) in enumerate(zip(headers, fills), 2):
            ws.cell(row=2, column=col).value = header
            ws.cell(row=2, column=col).fill = fill
            ws.cell(row=3, column=col).fill = fill
        
        # Добавляем заголовок фактических данных
        current_row = 5
        ws.cell(row=current_row, column=2).value = "Фактические данные"
        ws.cell(row=current_row, column=2).font = Font(bold=True)
        ws.merge_cells(start_row=current_row, start_column=2, end_row=current_row, end_column=9)
        
        # Заголовки для фактических данных
        current_row += 1
        for col, header in enumerate(headers, 2):
            ws.cell(row=current_row, column=col).value = header
            ws.cell(row=current_row, column=col).fill = header_fill
        
        # Получаем все сообщения и их анализы
        current_row += 1
        
        # Получаем список всех сообщений
        messages = self.get_all_messages()
        
        # Анализируем каждое сообщение
        for sender_name, date_str, message_text in messages:
            # Получаем результаты анализа из кэша
            cached_data = self.load_cached_analysis(message_text)
            
            if cached_data:
                # Если в сообщении есть несколько операций, обрабатываем каждую
                operations = []
                if "operations" in cached_data and cached_data["operations"]:
                    operations = cached_data["operations"]
                else:
                    # Если нет списка операций, создаем одну операцию из основных данных
                    operation = {
                        "work_type": cached_data.get("work_type", ""),
                        "operation": cached_data.get("operation", cached_data.get("work_type", "")),
                        "culture_from": cached_data.get("culture_from", ""),
                        "culture_to": cached_data.get("culture_to", ""),
                        "pu_number": cached_data.get("pu_number", ""),
                        "pu_area": cached_data.get("pu_area", ""),
                        "department": cached_data.get("department", ""),
                        "department_number": cached_data.get("department_number", ""),
                        "department_area": cached_data.get("department_area", ""),
                        "date": cached_data.get("date", date_str),
                        "subdivision": cached_data.get("subdivision", "АОР"),
                        "val_day": cached_data.get("val_day", ""),
                        "val_total": cached_data.get("val_total", "")
                    }
                    operations = [operation]
                
                # Для каждой операции добавляем строку в Excel
                for op in operations:
                    # Преобразуем культуру в полное название из справочника
                    culture = self.get_culture_name(op.get("culture_to", ""))
                    operation_name = self.get_operation_name(op.get("operation", op.get("work_type", "")))
                    
                    # Определяем дату
                    operation_date = op.get("date", "")
                    if not operation_date and "date_processed" in cached_data:
                        operation_date = cached_data["date_processed"]
                    elif not operation_date:
                        operation_date = date_str
                    
                    # Заполняем строку данными
                    ws.cell(row=current_row, column=2).value = operation_date
                    ws.cell(row=current_row, column=3).value = op.get("subdivision", "АОР")
                    ws.cell(row=current_row, column=4).value = operation_name
                    ws.cell(row=current_row, column=5).value = culture
                    ws.cell(row=current_row, column=6).value = op.get("pu_area", "")
                    ws.cell(row=current_row, column=7).value = op.get("total_area", op.get("pu_area", ""))
                    ws.cell(row=current_row, column=8).value = op.get("val_day", "")
                    ws.cell(row=current_row, column=9).value = op.get("val_total", "")
                    
                    current_row += 1
        
        # Настраиваем ширину колонок
        for col in range(1, 10):
            if col == 1:
                ws.column_dimensions[get_column_letter(col)].width = 40
            else:
                ws.column_dimensions[get_column_letter(col)].width = 20
        
        # Сохраняем файл в папку data/excel с названием, включающим дату и время
        now = datetime.now()
        date_str = now.strftime("%d.%m.%Y_%H-%M-%S")
        excel_dir = os.path.join('data', 'excel')
        os.makedirs(excel_dir, exist_ok=True)
        excel_path = os.path.join(excel_dir, f"Отчет_агрономов_{date_str}.xlsx")
        
        wb.save(excel_path)
        
        # Обновляем последний путь в статистике
        self.update_stats_with_excel(excel_path)
        
        return excel_path
        
    def add_excel_row_with_format(self, ws, row: int, date: str, subdivision: str, operation: str, 
                        culture: str, day_area: str, total_area: str, day_val: str = "", 
                        total_val: str = ""):
        """Добавляет строку в таблицу Excel с форматированием"""
        # Массив колонок (начиная с B, так как A зарезервирована для примеров)
        columns = list(range(2, 10))  # B to I
        
        # Добавляем значения в ячейки
        values = [date, subdivision, operation, culture, day_area, total_area, day_val, total_val]
        
        for col, value in zip(columns, values):
            cell = ws.cell(row=row, column=col)
            cell.value = value
            
            # Числовое форматирование для числовых данных
            if col >= 6:  # Для колонок E-I (числовые данные)
                try:
                    if value and str(value).strip():
                        cell.value = float(str(value).replace(",", ".").replace(" ", ""))
                        cell.number_format = '#,##0'
                except (ValueError, AttributeError):
                    pass
    
    def normalize_culture_abbreviation(self, culture: str) -> str:
        """Нормализует сокращения культур"""
        if not culture:
            return ""
            
        culture_lower = str(culture).lower().strip()
        
        # Прямые соответствия сокращений
        direct_map = {
            "оз": "Пшеница озимая товарная",
            "оз.": "Пшеница озимая товарная",
            "озимая": "Пшеница озимая товарная",
            "озим": "Пшеница озимая товарная",
            "пш": "Пшеница озимая товарная",
            "оз пш": "Пшеница озимая товарная",
            "мн": "Многолетние травы прошлых лет",
            "мн.": "Многолетние травы прошлых лет",
            "мн тр": "Многолетние травы прошлых лет",
            "многолетн": "Многолетние травы прошлых лет",
            "сах св": "Свекла сахарная",
            "сах": "Свекла сахарная",
            "св": "Свекла сахарная",
            "свекла": "Свекла сахарная",
            "соя": "Соя товарная",
            "сои": "Соя товарная",
            "зяби": "Зябь",
            "зябь": "Зябь",
            "кук": "Кукуруза товарная",
            "подс": "Подсолнечник товарный",
            "культ": "Пшеница озимая товарная"
        }
        
        # Проверка на прямое соответствие
        if culture_lower in direct_map:
            return direct_map[culture_lower]
        
        # Проверка на наличие словосочетаний
        for key, value in direct_map.items():
            if key in culture_lower:
                return value
        
        # Используем имеющийся словарь сокращений
        for abbr, full_name in self.reference_data["culture_abbreviations"].items():
            if abbr in culture_lower:
                return full_name
        
        # Если полное название в справочнике
        for full_name in self.reference_data["cultures"]:
            if culture_lower in full_name.lower() or full_name.lower() in culture_lower:
                return full_name
        
        # Логируем неизвестные культуры для дальнейшего анализа
        if len(culture) > 2:  # Чтобы не логировать пустые строки и короткие сокращения
            logging.info(f"Неизвестная культура: '{culture}'")
            
        return culture.capitalize()
    
    def get_statistics(self) -> Dict[str, Any]:
        """
        Возвращает статистику обработанных сообщений
        
        Returns:
            Dict[str, Any]: статистика сообщений
        """
        # Обновляем статистику перед возвратом
        stats = {
            "total_messages": self.statistics["total_messages"],
            "senders": self.statistics["senders"],
            "last_excel": self.current_excel_path
        }
        
        # Добавляем информацию о последнем Excel-файле
        if self.current_excel_path and os.path.exists(self.current_excel_path):
            stats["last_excel_modified"] = datetime.fromtimestamp(
                os.path.getmtime(self.current_excel_path)
            ).strftime("%d.%m.%Y %H:%M:%S")
        
        return stats 

    def clear_cache_and_history(self) -> None:
        """
        Очищает кэш анализа сообщений и историю для начала с чистого листа
        """
        try:
            # Очищаем кэш анализа
            cache_file = os.path.join(self.base_path, "analysis_cache.json")
            if os.path.exists(cache_file):
                os.remove(cache_file)
                logging.info(f"Кэш анализа удален: {cache_file}")
            
            # Сбрасываем счетчики и статистику
            self.message_counters = {}
            self.statistics = {
                "total_messages": 0,
                "senders": {},
                "fields_data": {}
            }
            
            logging.info("Кэш и счетчики очищены")
            return True
        except Exception as e:
            logging.error(f"Ошибка при очистке кэша: {str(e)}")
            return False 

    def update_stats_with_excel(self, excel_path: str) -> None:
        """Обновляет статистику с информацией о последнем Excel-файле"""
        self.current_excel_path = excel_path
        
    def load_cached_analysis(self, message_text: str) -> Dict[str, Any]:
        """
        Загружает результаты анализа сообщения из кэша
        
        Args:
            message_text (str): текст сообщения
            
        Returns:
            Dict[str, Any]: результаты анализа или пустой словарь
        """
        try:
            # Создаем уникальный ключ для сообщения
            cache_key = f"{len(message_text)}_{hash(message_text) % 10000}"
            
            # Пытаемся загрузить кэш
            cache_file = os.path.join(self.base_path, "analysis_cache.json")
            
            if os.path.exists(cache_file):
                with open(cache_file, 'r', encoding='utf-8') as f:
                    cache_data = json.load(f)
                    
                if cache_key in cache_data:
                    logging.info(f"Найден кэшированный анализ для сообщения")
                    return cache_data[cache_key]
            
            # Если кэш не найден, используем базовый парсинг
            result = self.parse_message(message_text)
            
            # Сохраняем результат в кэш для будущего использования
            self.save_to_cache(cache_key, result)
            
            return result
        except Exception as e:
            logging.error(f"Ошибка при загрузке кэша анализа: {str(e)}")
            return self.parse_message(message_text)
    
    def save_to_cache(self, cache_key: str, analysis_result: Dict[str, Any]) -> None:
        """
        Сохраняет результаты анализа в кэш
        
        Args:
            cache_key (str): ключ кэша
            analysis_result (Dict[str, Any]): результаты анализа
        """
        try:
            cache_file = os.path.join(self.base_path, "analysis_cache.json")
            
            # Загружаем существующий кэш или создаем новый
            if os.path.exists(cache_file):
                with open(cache_file, 'r', encoding='utf-8') as f:
                    cache_data = json.load(f)
            else:
                cache_data = {}
            
            # Добавляем новые данные
            cache_data[cache_key] = analysis_result
            
            # Сохраняем обновленный кэш
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(cache_data, f, ensure_ascii=False, indent=2)
                
            logging.info(f"Результаты анализа сохранены в кэш")
        except Exception as e:
            logging.error(f"Ошибка при сохранении в кэш: {str(e)}")
    
    async def analyze_and_cache_message(self, sender_name: str, message_text: str, yandex_gpt) -> Dict[str, Any]:
        """
        Анализирует сообщение с YandexGPT и сохраняет результаты в кэш
        
        Args:
            sender_name (str): имя отправителя
            message_text (str): текст сообщения
            yandex_gpt: экземпляр класса YandexGPT
            
        Returns:
            Dict[str, Any]: результаты анализа
        """
        # Создаем уникальный ключ для сообщения
        cache_key = f"{len(message_text)}_{hash(message_text) % 10000}"
        
        try:
            # Проверяем наличие кэша
            cache_file = os.path.join(self.base_path, "analysis_cache.json")
            
            if os.path.exists(cache_file):
                with open(cache_file, 'r', encoding='utf-8') as f:
                    cache_data = json.load(f)
                    
                if cache_key in cache_data:
                    logging.info(f"Используется кэшированный анализ")
                    return cache_data[cache_key]
            
            # Если в кэше нет, анализируем с YandexGPT
            logging.info(f"Анализируем сообщение с YandexGPT (не найдено в кэше)")
            result = await self.analyze_with_yandex_gpt(message_text, yandex_gpt)
            
            # Сохраняем результат в кэш
            self.save_to_cache(cache_key, result)
            
            return result
        except Exception as e:
            logging.error(f"Ошибка при анализе и кэшировании: {str(e)}")
            # Используем обычный парсинг при ошибке
            return self.parse_message(message_text)
    
    def get_culture_name(self, culture_abbr: str) -> str:
        """
        Получает полное название культуры из сокращения
        
        Args:
            culture_abbr (str): сокращение культуры
            
        Returns:
            str: полное название культуры
        """
        if not culture_abbr:
            return ""
        
        return self.normalize_culture_abbreviation(culture_abbr)
    
    def get_operation_name(self, operation_abbr: str) -> str:
        """
        Получает полное название операции из сокращения
        
        Args:
            operation_abbr (str): сокращение операции
            
        Returns:
            str: полное название операции
        """
        if not operation_abbr:
            return ""
        
        # Словарь типичных сокращений операций
        operation_map = {
            "пахота": "Пахота",
            "пах": "Пахота",
            "диск": "Дискование",
            "дискование": "Дискование",
            "диск 2": "Дискование 2-е",
            "диск 2-е": "Дискование 2-е",
            "культ": "Культивация",
            "предп культ": "Предпосевная культивация",
            "предпосевная": "Предпосевная культивация",
            "выравнивание": "Выравнивание",
            "подкормка": "Подкормка",
            "2-я подкормка": "2-я подкормка",
            "сев": "Сев",
            "посев": "Посев",
            "уборка": "Уборка"
        }
        
        # Проверяем прямое соответствие
        operation_lower = operation_abbr.lower()
        if operation_lower in operation_map:
            return operation_map[operation_lower]
        
        # Проверяем частичное соответствие
        for abbr, full_name in operation_map.items():
            if abbr in operation_lower:
                return full_name
        
        # Проверяем соответствие в списке стандартных операций
        for op in self.reference_data["operations"]:
            if op.lower() in operation_lower or operation_lower in op.lower():
                return op
        
        # Если не нашли совпадений, возвращаем исходное значение с заглавной буквы
        return operation_abbr.capitalize()
    
    def get_analysis_prompt(self, message_text: str) -> str:
        """
        Формирует промпт для анализа сообщения
        
        Args:
            message_text (str): Текст сообщения для анализа
            
        Returns:
            str: Сформированный промпт
        """
        return f"""Проанализируй сообщение или фотографию от агронома и извлеки структурированную информацию.

Если анализируется фотография:
1. Определи тип работы (вспашка, сев, уборка и т.д.)
2. Определи культуру
3. Оцени качество работы
4. Опиши состояние поля/культуры
5. Укажи любые проблемы или особенности, которые видны на фото

Если анализируется текстовое сообщение:
{message_text}

Верни результат в формате JSON со следующими полями:
{{
    "work_type": "тип работы",
    "operation": "конкретная операция",
    "culture_from": "исходная культура (если применимо)",
    "culture_to": "целевая культура (если применимо)", 
    "pu_number": "номер поля",
    "pu_area": "площадь в га",
    "department": "отделение",
    "quality": "оценка качества работы (если есть фото)",
    "field_condition": "состояние поля/культуры (если есть фото)",
    "issues": "обнаруженные проблемы (если есть)"
}}

Пример анализа текстового сообщения:
Входное сообщение: "Центральное отделение поле 125 площадь 82 га посев кукурузы"
{{
    "work_type": "посев",
    "operation": "посев",
    "culture_from": null,
    "culture_to": "кукуруза",
    "pu_number": "125",
    "pu_area": "82",
    "department": "Центральное отделение",
    "quality": null,
    "field_condition": null,
    "issues": null
}}

Пример анализа фото:
{{
    "work_type": "вспашка",
    "operation": "вспашка",
    "culture_from": null,
    "culture_to": null,
    "pu_number": null,
    "pu_area": null,
    "department": null,
    "quality": "хорошее",
    "field_condition": "почва хорошо обработана, без крупных комков",
    "issues": "небольшие пожнивные остатки на поверхности"
}}"""
    
    async def analyze_message(self, message_text: str, photo_path: str = None) -> dict:
        """
        Анализирует сообщение с помощью YandexGPT
        
        Args:
            message_text (str): Текст сообщения для анализа
            photo_path (str): Путь к фотографии (если есть)
            
        Returns:
            dict: Результат анализа
        """
        try:
            prompt = self.get_analysis_prompt(message_text)
            
            if photo_path:
                # Если есть фото, используем vision модель
                logging.info(f"Analyzing message with photo: {photo_path}")
                response = await self.yandex_gpt.generate_response(
                    prompt=prompt,
                    model="vision",
                    image_path=photo_path
                )
            else:
                # Если нет фото, используем обычную модель
                logging.info("Analyzing text message")
                response = await self.yandex_gpt.generate_response(
                    prompt=prompt,
                    model="yandexgpt"
                )

            if "error" in response:
                logging.error(f"Error from YandexGPT: {response['error']}")
                return {"error": response["error"]}

            result = self.yandex_gpt.get_response_text(response)
            
            try:
                # Пытаемся распарсить JSON из ответа
                parsed_result = json.loads(result)
                logging.info(f"Successfully parsed analysis result: {parsed_result}")
                return parsed_result
            except json.JSONDecodeError as e:
                logging.error(f"Failed to parse JSON from response: {str(e)}")
                return {"error": f"Failed to parse response: {str(e)}"}
                
        except Exception as e:
            logging.error(f"Error analyzing message: {str(e)}")
            return {"error": str(e)} 